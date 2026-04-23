#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成产业金融发展策略分析报告 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/Users/bismarck/KnowledgeBase/CreditWiki/outputs/产业金融发展策略分析报告.docx"

# 颜色
STONE_HEADING = RGBColor(0x3D, 0x40, 0x5B)
STONE_ACCENT = RGBColor(0x81, 0xB2, 0x9A)
STONE_BODY = RGBColor(0x1A, 0x1A, 0x2E)
STONE_MUTED = RGBColor(0xA8, 0xA8, 0xA8)
RED = RGBColor(0xCC, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

def set_font(run, name="微软雅黑", size=11, color=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold

def para_gaps(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

import docx.enum.text

def sh(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    for run in p.runs:
        set_font(run, size=16, color=STONE_HEADING, bold=True)

def subh(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_font(run, size=13, color=STONE_HEADING, bold=True)

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(16)
    for run in p.runs:
        set_font(run, size=11, color=STONE_BODY)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        set_font(run, size=11, color=STONE_BODY)
        run.text = text
    return p

def make_table(doc, headers, rows, col_widths_cm):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    for j, h in enumerate(headers):
        hdr.cells[j].text = h
        for p in hdr.cells[j].paragraphs:
            for run in p.runs:
                set_font(run, size=10, color=WHITE, bold=True)
        Shd = OxmlElement('w:shd')
        Shd.set(qn('w:fill'), "3D405B")
        hdr.cells[j]._tc.get_or_add_tcPr().append(Shd)
    # Data rows
    for i, row_data in enumerate(rows):
        row = tbl.rows[i+1]
        for j, text in enumerate(row_data):
            row.cells[j].text = text
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    set_font(run, size=10, color=STONE_BODY)
    # Col widths
    for row in tbl.rows:
        for j, w in enumerate(col_widths_cm):
            row.cells[j].width = Cm(w)
    return tbl

# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("【内部资料 · 仅供银行业内使用】")
set_font(r, size=9, color=RED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_gaps(p, before=20, after=6)
r = p.add_run("产业金融发展策略分析报告")
set_font(r, size=28, color=STONE_HEADING, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_gaps(p, after=20)
r = p.add_run("——基于龙岩分行汇报的第一性原理分析")
set_font(r, size=16, color=STONE_ACCENT)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_gaps(p, after=20)
r = p.add_run("─" * 40)
set_font(r, size=10, color=STONE_MUTED)

from datetime import datetime
meta_items = [
    "分析对象：福建省龙岩市\u201c2+4\u201d产业体系",
    "对标借鉴：上海生物医药产业金融服务方案",
    f"报告编制：Hermes Agent",
    f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
]
for m in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_gaps(p, before=3, after=3)
    r = p.add_run(m)
    set_font(r, size=12, color=STONE_BODY)

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_gaps(p, before=40)
r = p.add_run("CreditWiki 知识库 | 兴业银行研究团队")
set_font(r, size=10, color=STONE_MUTED)

add_page_break(doc)

# ============================================================
# 目录
# ============================================================
p = doc.add_paragraph()
para_gaps(p, after=10)
r = p.add_run("目  录")
set_font(r, size=18, color=STONE_HEADING, bold=True)

toc = [
    ("执行摘要", "3"),
    ("第一部分：龙岩分行汇报要点分析", "4"),
    ("第二部分：第一性原理拆解", "6"),
    ("第三部分：上海地区产业分析", "8"),
    ("第四部分：上海地区产业金融布局建议", "10"),
    ("结论", "13"),
]
for item, pg in toc:
    p = doc.add_paragraph()
    para_gaps(p, before=4, after=4)
    r = p.add_run(item + "  " + "\u2022" * 3 + "  " + pg)
    set_font(r, size=11, color=STONE_BODY)

add_page_break(doc)

# ============================================================
# 执行摘要
# ============================================================
sh(doc, "执行摘要")

summaries = [
    "本报告基于龙岩分行《产业金融发展思路汇报》（练文锋，2026年4月）以及兴业银行上海分行《生物医药产业金融服务方案》，运用第一性原理方法论，对产业金融的核心逻辑进行拆解，并就上海地区产业金融布局提出系统性建议。",
    "核心发现：龙岩分行的成功经验可归结为\u201c腾笼换鸟\u201d战略——三年压降36.2亿元地产与平台融资，置换出68亿元产业贷款，21条产业链客户贷款余额183亿元、占比达88%。这一打法揭示了产业金融的本质规律：资金应追随产业现金流，而非抵押物。",
    "上海与龙岩代表两种截然不同的产业金融范式：龙岩是\u201c资源型\u201d，核心痛点在账期错配与价格波动；上海是\u201c创新型\u201d，核心痛点在研发阶段资金饥渴与轻资产矛盾。金融工具适配逻辑因此根本不同。",
    "核心建议：上海产业金融应采用\u201c以投行思维做商行业务\u201d的策略——用股权收益覆盖研发风险，用供应链金融锁定产业现金流，用临港/张江等政策红利区做批量获客入口。",
]
for s in summaries:
    body(doc, s)

p = doc.add_paragraph()
para_gaps(p, before=10, after=4)
r = p.add_run("核心数据速览")
set_font(r, size=12, color=STONE_ACCENT, bold=True)

make_table(doc,
    ["指标", "数值"],
    [
        ["龙岩六大产业贷款余额", "68亿元（2025年末）"],
        ["十五五目标（2030年）", "500亿元（翻一番）"],
        ["上海生物医药产业规模", "突破1万亿元（2025年）"],
        ["张江药谷企业数量", "2300+家"],
    ],
    [7, 10]
)
p = doc.add_paragraph()
r = p.add_run("表1：核心数据速览")
set_font(r, size=9, color=STONE_MUTED)

add_page_break(doc)

# ============================================================
# 第一部分
# ============================================================
sh(doc, "第一部分：龙岩分行汇报要点分析")

subh(doc, "一、龙岩\u201c2+4\u201d产业体系")
body(doc, "龙岩市\u201c十五五\u201d期间工业产业体系由2大支柱产业和4大战略性新兴产业构成：")

make_table(doc,
    ["类别", "产业名称", "2025年产值（亿元）", "我行授信（亿元）", "用信占比"],
    [
        ["支柱产业", "有色金属", "1512", "162", "10%"],
        ["支柱产业", "机械装备", "920", "553", "29%"],
        ["新兴产业", "新材料", "280", "71", "18%"],
        ["新兴产业", "新能源", "240", "47", "20%"],
        ["新兴产业", "电子信息", "150", "34", "25%"],
        ["新兴产业", "节能环保", "20", "13", "64%"],
        ["合计", "—", "3122", "880", "—"],
    ],
    [2.5, 2.8, 3.5, 3.5, 2.5]
)
p = doc.add_paragraph()
r = p.add_run("表2：龙岩市\u201c2+4\u201d产业体系核心数据（来源：龙岩分行汇报）")
set_font(r, size=9, color=STONE_MUTED)

subh(doc, "二、核心打法：腾笼换鸟")
body(doc, "龙岩分行的战略核心在于主动压降高风险资产，置换为产业资产：")
for b in [
    "三年压降房地产融资7.4亿元、政府平台及隐债项目28.8亿元，合计36.2亿元",
    "六大产业贷款余额从45.6亿元增长至68亿元，三年新增22.4亿元",
    "21条重点产业链客户497户，贷款余额183亿元，三年累计新增45亿元",
    "区域前十大行业贷款偏离度仅3.5%，省内排名第二",
]:
    bullet(doc, b)

subh(doc, "三、标杆案例：上杭金铜产业链")
body(doc, "分行打造了可复制的产业链金融服务标杆——上杭金铜及新材料产业集群：")
make_table(doc,
    ["维度", "数值"],
    [
        ["服务客户数", "59户（含金铜产业链40户、新材料26户）"],
        ["累计授信", "60亿元"],
        ["开户合作率", "近90%"],
        ["授信覆盖率", "超62%"],
        ["核心方法", "\u201c服务一个龙头、带动一条链条、激活一片生态\u201d"],
    ],
    [3, 12]
)

subh(doc, "四、十五五目标规划")
body(doc, "分行制定了明确的量化目标：")
for b in [
    "2025-2030年，六大主导产业用信占比从33%提升至50%（对应绝对额翻番）",
    "2030年末，六大主导产业用信金额突破500亿元，较2025年末翻一番",
    "重点产业融资增速持续高于各项企业贷款平均增速、高于当地同业平均增速",
]:
    bullet(doc, b)

add_page_break(doc)

# ============================================================
# 第二部分
# ============================================================
sh(doc, "第二部分：第一性原理拆解")

subh(doc, "一、产业金融的本质规律")
body(doc, "从第一性原理出发，金融的核心功能是跨时空配置资金。产业金融的根本逻辑在于：识别产业价值链中谁在占用谁的资金，由此形成谁的负债，银行服务那个\u201c资金缺口节点\u201d。")

p = doc.add_paragraph()
para_gaps(p, before=8, after=4)
r = p.add_run("资金缺口节点分析框架：")
set_font(r, size=12, color=STONE_HEADING, bold=True)
body(doc, "任何产业都可以用这个框架拆解：原料采购占压资金（上游）\u2192 生产加工占压资金（中游）\u2192 销售回款慢（下游）\u2192 谁最缺钱，谁就是银行的客户。")

subh(doc, "二、龙岩模式与上海模式的本质差异")
make_table(doc,
    ["维度", "龙岩（资源型）", "上海（创新型）"],
    [
        ["代表产业", "有色金属、机械装备", "生物医药、集成电路"],
        ["资金缺口节点", "采掘/加工环节占压", "研发阶段最缺钱"],
        ["核心风险来源", "大宗商品价格周期", "研发失败/政策变更"],
        ["核心抓手", "核心企业信用穿透", "管线估值/订单确权"],
        ["适配金融工具", "供应链票据+保理", "投贷联动+专利质押"],
        ["盈利模型", "利差为主+结算沉淀", "利息+股权收益+托管"],
        ["风控逻辑", "盯住抵押物+核心企业", "盯住技术里程碑+估值"],
    ],
    [3, 6.5, 6.5]
)
p = doc.add_paragraph()
r = p.add_run("表3：龙岩模式与上海模式第一性原理对比")
set_font(r, size=9, color=STONE_MUTED)

subh(doc, "三、产业金融的成功公式")
body(doc, "综合龙岩经验，产业金融成功可以提炼为以下公式：")
p = doc.add_paragraph()
para_gaps(p, before=8, after=4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("产业金融成功 = 锁定产业链核心企业 \u00d7 全链条渗透 \u00d7 综合收益覆盖风险")
set_font(r, size=14, color=STONE_ACCENT, bold=True)
body(doc, "三个乘数项缺一不可：")
for b in [
    "锁定核心企业：找到产业链中议价能力最强、信用最稳定的\u201c链主\u201d企业",
    "全链条渗透：以核心企业为圆心，沿着供应链反向追踪上下游中小客户",
    "综合收益覆盖风险：用利息收入+结算沉淀+托管收入+潜在股权收益综合评估项目回报",
]:
    bullet(doc, b)

add_page_break(doc)

# ============================================================
# 第三部分
# ============================================================
sh(doc, "第三部分：上海地区产业分析")

subh(doc, "一、上海生物医药产业全景")
body(doc, "上海生物医药产业2025年正式迈入万亿级产业集群，成为中国生物医药产业高地。")
make_table(doc,
    ["维度", "数据"],
    [
        ["2025年产业规模", "突破1万亿元"],
        ["2024年产业规模", "9847亿元（较上年+510亿元）"],
        ["年复合增长率", "8.94%"],
        ["生物医药企业数量", "超过4000家"],
        ["科创板上市企业数量", "全国第一"],
        ["张江药谷企业数", "2300+家"],
        ["临港生命蓝湾企业数", "100+家"],
        ["全国生物医药人才占比", "约1/3"],
    ],
    [5, 10]
)

subh(doc, "二、产业链结构与资金缺口分析")
body(doc, "上海生物医药产业链可细分为以下环节，各环节资金缺口特征各异：")
make_table(doc,
    ["产业链环节", "代表企业", "核心痛点", "适配金融工具"],
    [
        ["创新药物研发", "君实生物、映恩生物", "研发周期长（5-10年）、轻资产、商业化前无现金流", "投贷联动、BD交易融资"],
        ["CRO（研发外包）", "药明康德、美迪西", "产能建设投入大、客户质量为王", "供应链金融、授信池"],
        ["CDMO（生产外包）", "和元生物、凯莱英", "重资产、产能利用率关键", "设备融资租赁、并购贷款"],
        ["高端医疗器械", "联影医疗、微创机器人", "技术壁垒高、国产替代空间大", "买方信贷、设备融资租赁"],
        ["细胞与基因治疗", "复星凯特、药明巨诺", "前沿赛道、监管创新、产业化初期", "专项产业基金+银行授信联动"],
    ],
    [2.8, 3.5, 5.5, 4]
)

subh(doc, "三、核心企业画像与金融需求")
body(doc, "创新药企领域（高风险高回报，适合投贷联动）：")
for b in [
    "君实生物：特瑞普利单抗中美欧三地获批，商业化扩张+研发投入，BD交易支撑",
    "映恩生物：12款ADC候选药物，港股上市，IPO融资+临床推进并重",
    "英矽智能：AI平台发现27款临床前候选化合物，港股上市（AI制药第一股）",
]:
    bullet(doc, b)

body(doc, "CRO/CDMO领域（现金流相对稳定，适合供应链金融）：")
for b in [
    "药明康德：全球TOP3 CRO，6000+活跃客户，在手订单493亿元",
    "泰格医药：国内临床CRO龙头，全球化布局",
]:
    bullet(doc, b)

subh(doc, "四、政策红利区分析")
make_table(doc,
    ["政策区域", "核心政策红利", "适合业务"],
    [
        ["临港新片区", "细胞与基因治疗创新监管政策、国际多地准入", "跨境并购贷款、CGT专项融资"],
        ["张江药谷", "创新药物绿色通道、首批国产替代支持", "投贷联动、专利质押融资"],
        ["外高桥保税区", "离岸贸易、国际贸易融资便利", "离岸保理、国际信用证"],
        ["东方美谷", "美丽健康产业注册人制度", "消费金融、产业链供应链金融"],
    ],
    [3, 7, 6]
)

add_page_break(doc)

# ============================================================
# 第四部分
# ============================================================
sh(doc, "第四部分：上海地区产业金融布局建议")

subh(doc, "一、总体策略：以投行思维做商行业务")
body(doc, "上海产业金融与龙岩的资源型打法有本质区别，必须采用\u201c以投行思维做商行业务\u201d的策略：")
for b in [
    "用股权收益覆盖研发风险（明股实债/远期认股权）",
    "用供应链金融锁定产业现金流（应收账款保理/订单融资）",
    "用政策红利区做批量获客入口（临港/张江/外高桥）",
]:
    bullet(doc, b)

subh(doc, "二、分片区布局策略")
make_table(doc,
    ["片区", "核心策略", "重点客户", "主打产品"],
    [
        ["张江药谷", "锁定ADC/GLP-1/AI制药前沿赛道", "君实生物、映恩生物、英矽智能", "投贷联动、明股实债"],
        ["临港生命蓝湾", "CGT+生物制品创新监管试点", "复星凯特、药明巨诺、和元生物", "跨境并购贷款、专项产业基金"],
        ["外高桥保税区", "离岸贸易+国际贸易融资", "全球顶尖药企中国总部、进口商", "离岸保理、国际信用证、福费廷"],
        ["CRO/CDMO集群", "锁定现金流稳定的头部客户", "药明康德、泰格医药、美迪西", "供应链票据、授信池、现金管理"],
    ],
    [2.5, 4, 4.5, 4.8]
)

subh(doc, "三、产品创新方向")
body(doc, "1. 科创金融产品（针对轻资产研发企业）：")
for b in [
    "管线估值贷款：以药品管线（如临床II期、III期）估值作为第二还款来源",
    "订单确认书融资：锁定商业化早期的药品订单，提前提供流动性支持",
    "专利技术质押：在传统抵押基础上叠加专利、软件著作权质押增信",
]:
    bullet(doc, b)

body(doc, "2. 供应链金融产品（针对产业链上下游）：")
for b in [
    "反向保理：与核心药企确权，让上游原料商提前收回应收账款",
    "商业票据：推动核心药企开立供应链票据，中小供应商持票贴现",
]:
    bullet(doc, b)

body(doc, "3. 跨境金融产品（针对国际化药企）：")
for b in [
    "跨境并购贷款：支持中国药企海外收购药品权益、CDMO产能",
    "BD交易融资：支持创新药海外License-out的里程碑付款节点",
]:
    bullet(doc, b)

subh(doc, "四、复制\u201c龙岩上杭模式\u201d到上海")
body(doc, "龙岩上杭模式的本质是：锁定1个链主龙头企业 \u2192 穿透上下游中小客户 \u2192 实现开户+授信+结算全覆盖。这套打法完全可复制到上海生物医药产业链：")
make_table(doc,
    ["步骤", "龙岩上杭模式", "上海张江模式"],
    [
        ["锁定龙头", "上杭金铜产业链链主", "张江药谷复宏汉霖/君实生物"],
        ["穿透链条", "59户上下游企业", "预计200+家CRO/CDMO/原料供应商"],
        ["开户目标", "近90%", "85%以上"],
        ["授信覆盖", "超62%", "目标60%以上"],
    ],
    [2.5, 6.5, 7.5]
)

subh(doc, "五、落地执行建议")
make_table(doc,
    ["优先级", "举措", "参考龙岩模式", "预计效果"],
    [
        ["P0", "张江药谷生物医药专班（研究+营销+审批闭环）", "四位一体专班体系", "2026年落地首批50家客户"],
        ["P0", "临港生命蓝湾跨境并购贷款试点", "并购贷款跨境经验", "年内投放10亿元"],
        ["P1", "复制\u201c上杭模式\u201d到张江产业链", "锁定1龙头\u2192穿透全链", "新增对公客户100+户"],
        ["P1", "搭建生物医药产业数字图谱", "龙岩14条产业链图谱", "精准画像+批量获客"],
        ["P2", "科创金融产品创新（专利质押+管线融资）", "四懂两会产业铁军", "差异化竞争护城河"],
    ],
    [1.5, 5.5, 4.5, 4.5]
)

add_page_break(doc)

# ============================================================
# 结论
# ============================================================
sh(doc, "结  论")

conclusions = [
    "龙岩分行的\u201c腾笼换鸟\u201d战略揭示了产业金融的第一性原理：资金应追随产业现金流，而非抵押物。龙岩三年压降36.2亿元高风险资产、五年内实现产业贷款翻番的经验，证明了主动的结构调整比被动的风险处置有更高的长期回报。",
    "上海与龙岩代表两种截然不同的产业金融范式，但底层逻辑一致：识别产业价值链中的资金缺口节点，用适配的金融工具服务那个节点，并通过全链条渗透放大收益、分散风险。",
    "上海生物医药产业已进入万亿级集群时代，张江药谷、临港生命蓝湾、外高桥保税区等政策红利区为银行提供了天然的批量获客入口。以投行思维做商行业务，用股权收益覆盖研发风险，是上海产业金融的核心破局点。",
    "建议上海分行尽快启动生物医药产业专班建设，借鉴龙岩\u201c四懂两会\u201d人才体系和\u201c四位一体\u201d专班机制，优先在张江药谷和临港新片区建立标杆项目，形成可复制的上海模式。",
]

for i, c in enumerate(conclusions):
    p = doc.add_paragraph()
    para_gaps(p, before=8, after=8)
    r1 = p.add_run(f"【结论{i+1}】")
    set_font(r1, size=12, color=STONE_ACCENT, bold=True)
    r2 = p.add_run(c)
    set_font(r2, size=11, color=STONE_BODY)
    p.paragraph_format.line_spacing = Pt(18)

p = doc.add_paragraph()
para_gaps(p, before=30)
r = p.add_run("\u2500" * 40)
set_font(r, size=10, color=STONE_MUTED)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_gaps(p, before=4)
r = p.add_run("本报告由 CreditWiki 知识库 AI 分析生成 | 仅供参考，不构成投资建议")
set_font(r, size=9, color=STONE_MUTED)

doc.save(OUTPUT)
print(f"\u2705 报告已生成：{OUTPUT}")
